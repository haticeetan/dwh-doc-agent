import io
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _apply_heading(doc: Document, text: str, level: int):
    """Başlık ekler."""
    doc.add_heading(text, level=level)


def _add_table_from_markdown(doc: Document, lines: list[str]):
    """
    Markdown tablo satırlarını Word tablosuna dönüştürür.
    Örnek:
        | Kolon | Tip | Açıklama |
        |-------|-----|----------|
        | ID    | NUM | ...      |
    """
    rows = [l for l in lines if l.strip().startswith("|") and "---" not in l]
    if not rows:
        return

    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        parsed.append(cells)

    if not parsed:
        return

    col_count = len(parsed[0])
    table = doc.add_table(rows=len(parsed), cols=col_count)
    table.style = "Table Grid"

    for r_idx, row in enumerate(parsed):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = cell_text
            # İlk satır başlık — kalın yap
            if r_idx == 0:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                cell._tc.get_or_add_tcPr()
                # Başlık hücre arka planı
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "2E74B5")
                shading.set(qn("w:color"), "auto")
                shading.set(qn("w:val"), "clear")
                cell._tc.get_or_add_tcPr().append(shading)

    doc.add_paragraph()


def _parse_and_render(doc: Document, markdown: str):
    """
    Markdown metni satır satır okur, Word elementlerine dönüştürür.
    Desteklenen formatlar: # başlıklar, ** kalın, tablolar, normal paragraf.
    """
    lines = markdown.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Başlıklar
        if line.startswith("#### "):
            _apply_heading(doc, line[5:].strip(), level=4)
        elif line.startswith("### "):
            _apply_heading(doc, line[4:].strip(), level=3)
        elif line.startswith("## "):
            _apply_heading(doc, line[3:].strip(), level=2)
        elif line.startswith("# "):
            _apply_heading(doc, line[2:].strip(), level=1)

        # Markdown tablo — | ile başlayan satır bloğunu topla
        elif line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            _add_table_from_markdown(doc, table_lines)
            continue

        # Yatay çizgi
        elif line.strip() in ("---", "***", "___"):
            doc.add_paragraph("─" * 60)

        # Boş satır
        elif line.strip() == "":
            pass

        # Normal paragraf — **kalın** işle
        else:
            para = doc.add_paragraph()
            parts = re.split(r"(\*\*[^*]+\*\*)", line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                else:
                    para.add_run(part)

        i += 1


def build_docx(markdown_text: str, title: str = "Tablo Dokümantasyonu") -> bytes:
    """
    Markdown metni alır, Word belgesi üretir, bytes olarak döner.

    Args:
        markdown_text: doc_writer'ın ürettiği markdown içerik
        title:         Belge başlığı (tablo adı)

    Returns:
        .docx dosyasının byte içeriği
    """
    doc = Document()

    # Sayfa kenar boşlukları
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Üst bilgi
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = f"DW-DocAgent | {title}"
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # İçeriği işle
    _parse_and_render(doc, markdown_text)

    # Bytes'a çevir
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()